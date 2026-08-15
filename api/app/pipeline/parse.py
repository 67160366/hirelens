"""Turn an uploaded document (PDF or DOCX) into text with stable character offsets.

Two things make this more than a wrapper around `extract_text()`:

1.  Offsets are the contract the rest of the system depends on. `ParsedDocument.text`
    is the single coordinate space that evidence spans point into, and page
    boundaries are recorded so a span can be mapped back to a page for the UI.
2.  Text is NFC-normalized per page *before* joining. Normalizing later would shift
    every offset computed before it, which silently corrupts Thai combining marks.

A page with no text layer is handed to OCR when an engine is supplied, and the
recognized text is substituted into the page list *before* spans are measured — so
a rescued page is indistinguishable from a normal one to everything downstream.
Without an engine the page is reported rather than quietly yielding an empty
document.
"""

from __future__ import annotations

import io
import logging
import unicodedata
from bisect import bisect_right
from collections.abc import Iterator, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Self

import docx
import docx.document
import pdfplumber
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.pipeline import geometry as geometry_module
from app.pipeline.geometry import CharRun, PageGeometry, runs_for
from app.pipeline.layout import detect_reading_order, extract_in_reading_order
from app.pipeline.ocr import OCREngine, OCRError, OCRUnavailableError

# Resumes are PII: log page numbers and error types, never text.
logger = logging.getLogger(__name__)

# A page with fewer than this many non-whitespace characters is treated as having
# no text layer. Scanned pages usually extract to nothing at all, but a stray
# ligature or page number can leak through.
MIN_CHARS_PER_TEXT_PAGE = 20

PAGE_SEPARATOR = "\n\n"


def _needs_ocr(text: str) -> bool:
    """Whether a page yielded too little text to count as having a text layer.

    One rule, used both to pick the pages OCR should run on and to decide what
    `pages_without_text` reports afterwards — so the two can never disagree.
    """
    return len(text.strip()) < MIN_CHARS_PER_TEXT_PAGE


class ParseError(Exception):
    """Base class for document parsing failures."""


class UnsupportedFileTypeError(ParseError):
    def __init__(self, suffix: str) -> None:
        super().__init__(f"Unsupported file type: {suffix or '(no extension)'}")
        self.suffix = suffix


class CorruptDocumentError(ParseError):
    """The file could not be opened as the type its extension claims."""


class NoTextLayerError(ParseError):
    """Every page is an image, and OCR either was not available or found nothing."""

    def __init__(self, page_count: int, *, ocr_attempted: bool = False) -> None:
        if ocr_attempted:
            detail = "OCR ran but recognized no usable text; the scan may be too low quality."
        else:
            detail = "It is a scan and requires OCR, which is not enabled."
        super().__init__(
            f"Document has no text layer on any of its {page_count} page(s) but does "
            f"contain images. {detail}"
        )
        self.page_count = page_count
        self.ocr_attempted = ocr_attempted


class EmptyDocumentError(ParseError):
    """No text and no images. Distinct from a scan on purpose.

    A scan is recoverable by OCR; a blank document is not. Collapsing the two would
    send blank pages through the OCR path in M2 and report a misleading reason to
    the user.
    """

    def __init__(self, page_count: int) -> None:
        super().__init__(
            f"Document has no text and no images across its {page_count} page(s); "
            "it appears to be blank."
        )
        self.page_count = page_count


@dataclass(frozen=True, slots=True)
class PageSpan:
    """Where one page's text sits inside `ParsedDocument.text`."""

    page_number: int
    """1-indexed, matching what a person sees in a PDF viewer."""

    char_start: int
    char_end: int

    @property
    def is_empty(self) -> bool:
        return self.char_end <= self.char_start


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    text: str
    """NFC-normalized full text. The coordinate space for all evidence offsets."""

    pages: tuple[PageSpan, ...]
    pages_without_text: tuple[int, ...]
    """1-indexed pages that *still* yielded no usable text, after OCR if it ran.

    Kept as the honest work list: a page OCR rescued is not in here, because it has
    text now."""

    pages_from_ocr: tuple[int, ...] = ()
    """1-indexed pages whose text was recognized from an image rather than read
    from a text layer. Surfaced to the user, because a citation into one of these
    is faithful to what was read, not necessarily to what was printed."""

    page_geometry: tuple[PageGeometry, ...] = ()
    """Where each character sits on the page it was read from (M5 slice 3).

    **Sparse on purpose, and a missing page is a supported state.** A page is absent
    when its geometry could not be proven consistent with its text, when OCR replaced
    the text wholesale, or when the document has no glyph boxes at all (`.docx`). The
    overlay falls back to the text pane for those and says why; a silent fallback is
    indistinguishable from a bug, and a *guessed* box is worse than either."""

    # Page start offsets, precomputed for O(log n) lookup: a resume yields dozens of
    # spans to map and page_for_offset runs on each one.
    _page_starts: tuple[int, ...] = field(init=False, repr=False, compare=False, default=())

    def __post_init__(self) -> None:
        object.__setattr__(self, "_page_starts", tuple(p.char_start for p in self.pages))

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def is_partially_scanned(self) -> bool:
        """Some pages had text and some did not — worth surfacing to the user."""
        return bool(self.pages_without_text) and len(self.pages_without_text) < self.page_count

    @property
    def used_ocr(self) -> bool:
        return bool(self.pages_from_ocr)

    def page_for_offset(self, offset: int) -> int:
        """Return the 1-indexed page containing `offset`.

        Offsets outside the text clamp to the first or last page rather than
        raising: a citation should never fail to render because of an off-by-one.
        """
        if not self.pages:
            return 1
        index = bisect_right(self._page_starts, offset) - 1
        return self.pages[max(index, 0)].page_number

    @property
    def stored_page_spans(self) -> list[dict[str, int]]:
        """The page boundaries in the shape `resumes.page_spans` holds.

        Written out field by field rather than via `asdict` so the stored key names
        are greppable, and kept beside `from_stored` because the two are the halves
        of one round-trip: neither may change without the other.
        """
        return [
            {
                "page_number": span.page_number,
                "char_start": span.char_start,
                "char_end": span.char_end,
            }
            for span in self.pages
        ]

    @property
    def stored_page_geometry(self) -> list[dict[str, Any]]:
        """The character geometry in the shape `resumes.page_geometry` holds.

        Beside `stored_page_spans` and paired with `from_stored` for the same reason:
        neither half of a round-trip may change without the other.
        """
        return geometry_module.stored(self.page_geometry)

    @classmethod
    def from_stored(
        cls,
        text: str,
        page_spans: list[dict[str, int]] | None,
        *,
        pages_without_text: Sequence[int] = (),
        pages_from_ocr: Sequence[int] = (),
        page_geometry: list[dict[str, Any]] | None = None,
    ) -> Self:
        """Rebuild a document from values already stored on a resume row.

        This is what lets a *new* quote — one located in stored text long after the
        upload, which is what judging does — be mapped back to a page. It re-reads
        nothing: `text` is the verbatim `document_text` that every evidence offset
        already indexes into, so no citation can shift.

        That is the whole difference from `reparse_document` in
        `services/resume_service.py`, which goes back to the stored *file*: pass it
        a different OCR configuration and a page rescued then but not now comes back
        empty, moving every offset after it. This function cannot do that, because
        it never parses anything.

        `page_spans` is None on rows written before migration `0005`. `pages` is
        then empty and `page_for_offset` answers 1 for every offset — the honest
        result, since those rows never recorded where their pages ended. Backfilling
        them would mean re-parsing under the identical OCR configuration, which is
        exactly the hazard above.
        """
        return cls(
            text=text,
            pages=tuple(PageSpan(**span) for span in page_spans or ()),
            pages_without_text=tuple(pages_without_text),
            pages_from_ocr=tuple(pages_from_ocr),
            page_geometry=geometry_module.from_stored(page_geometry),
        )


def parse_document(path: Path, *, ocr: OCREngine | None = None) -> ParsedDocument:
    """Parse a document by file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path, ocr=ocr)
    if suffix == ".docx":
        return parse_docx(path)
    raise UnsupportedFileTypeError(suffix)


def parse_document_bytes(
    data: bytes, *, filename: str, ocr: OCREngine | None = None
) -> ParsedDocument:
    """Parse an in-memory upload, dispatching on the filename's extension.

    Uploads arrive as bytes; going via a temp file would add I/O and a cleanup
    path for no benefit.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(io.BytesIO(data), ocr=ocr)
    if suffix == ".docx":
        return parse_docx(io.BytesIO(data))
    raise UnsupportedFileTypeError(suffix)


def parse_pdf(source: Path | io.BytesIO, *, ocr: OCREngine | None = None) -> ParsedDocument:
    """Read a PDF, optionally recovering text-less pages with OCR.

    `ocr=None` means OCR is off, which is deliberately not the same as an engine
    that ran and found nothing — the second is an answer about the document.
    """
    try:
        with pdfplumber.open(source) as pdf:
            raw_pages: list[str] = []
            # Whether a text-less page carries an image decides if OCR can rescue
            # it. Tracked per page so a blank page is never rendered and paid for.
            page_has_images: list[bool] = []
            # Where each character sits, per page, parallel to `raw_pages`. `None`
            # for a page whose geometry could not be trusted (M5 slice 3).
            raw_geometry: list[_PageInk | None] = []
            for page in pdf.pages:
                text, runs = _text_of(page)
                raw_pages.append(text)
                page_has_images.append(bool(page.images))
                raw_geometry.append(
                    None
                    if runs is None
                    else _PageInk(width=float(page.width), height=float(page.height), runs=runs)
                )

            from_ocr: tuple[int, ...] = ()
            if ocr is not None:
                from_ocr = _recover_pages_with_ocr(pdf.pages, raw_pages, page_has_images, ocr)
                for page_number in from_ocr:
                    # OCR replaces a page's text wholesale, so every offset the
                    # text layer's geometry described is gone. Dropping it is the
                    # honest answer and the one slice 4 already plans to render:
                    # a recognized page has no glyph boxes to overlay onto.
                    raw_geometry[page_number - 1] = None
    except (ParseError, OCRError):
        # An OCR fault is about the engine, not the document: letting it become a
        # CorruptDocumentError below would mark a fixable misconfiguration as a
        # permanently broken file.
        raise
    except Exception as exc:  # pdfplumber/pdfminer raise a wide variety of types
        raise CorruptDocumentError(f"Could not read PDF: {exc}") from exc

    return _assemble(
        raw_pages,
        has_images=any(page_has_images),
        pages_from_ocr=from_ocr,
        ocr_attempted=ocr is not None,
        ink=raw_geometry,
    )


@dataclass(frozen=True, slots=True)
class _PageInk:
    """One page's geometry before `_assemble` rebases it into document space.

    Its runs index into that page's *raw* text — before NFC and the NUL strip — so
    it is deliberately not `PageGeometry`, which is measured against the stored
    `document_text` and is the only shape anything outside this module sees.
    """

    width: float
    height: float
    runs: tuple[CharRun, ...]


def _text_of(page: Any) -> tuple[str, tuple[CharRun, ...] | None]:
    """Read one page, one column at a time when it has more than one.

    `detect_reading_order` answers `None` for every page it is not confident about,
    and that branch is the one that ran before column detection existed — so a
    single-column document parses to exactly the string it always did, and no
    citation already shown to a user can shift.

    Returns the character geometry beside the text (M5 slice 3). It is measured from
    the same object the text came from and checked against the text itself, so it is
    `None` rather than approximate whenever the two could disagree — see
    `pipeline/geometry.py`. **The text is produced by exactly the calls it always
    was**, which is what keeps `document_text` byte-identical.
    """
    boxes = detect_reading_order(page)
    if boxes is None:
        text: str = page.extract_text() or ""
        return text, runs_for(page, text)
    return extract_in_reading_order(page, boxes)


def parse_docx(source: Path | io.BytesIO) -> ParsedDocument:
    """Read a .docx into the same offset space a PDF produces.

    **A .docx has no pages.** Word reflows text at render time, so where page 2
    begins depends on the fonts, the printer and the zoom level — it is not in the
    file. Rather than invent page numbers, the whole document is reported as one
    page: a citation then says "somewhere in this document", which is true, instead
    of "page 2", which would be a guess dressed as a fact. Explicit author-inserted
    page breaks *are* in the file and could split it later; automatic ones never
    will be.

    Tables are read in document order along with paragraphs, because resumes
    routinely use a table for layout and skipping them would silently lose the
    skills section — the failure would look like a model that missed things.
    """
    try:
        # python-docx takes a path as `str` rather than `Path`; a BytesIO passes
        # straight through as the file object it already is.
        document = docx.Document(str(source) if isinstance(source, Path) else source)
        text = "\n".join(_iter_docx_text(document))
    except ParseError:
        raise
    except Exception as exc:  # python-docx/lxml raise a wide variety of types
        raise CorruptDocumentError(f"Could not read DOCX: {exc}") from exc

    # `has_images=False`: an image-only .docx is not a scan in the sense OCR could
    # rescue, and reporting it as one would send the user chasing a setting that
    # would not help.
    return _assemble([text])


def _iter_docx_text(document: docx.document.Document) -> Iterator[str]:
    """Yield paragraph and table text in the order the document declares it.

    `document.paragraphs` skips anything inside a table, so the body's own child
    order is walked instead — it is the only place that knows a table sits between
    two paragraphs rather than after all of them.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document).text
        elif child.tag == qn("w:tbl"):
            for row in Table(child, document).rows:
                # Tab between cells, newline between rows: it keeps a row that
                # reads as one line in Word reading as one line here, which is
                # what a quote from it will look like.
                yield "\t".join(cell.text.strip() for cell in row.cells)


def _recover_pages_with_ocr(
    pages: list[Any],
    raw_pages: list[str],
    page_has_images: list[bool],
    ocr: OCREngine,
) -> tuple[int, ...]:
    """Fill text-less pages from OCR, **editing `raw_pages` in place**.

    Returns the 1-indexed pages that OCR actually rescued. A page is only replaced
    when the recognized text clears the same bar a text layer has to clear, so a
    handful of noise characters never enters `document_text` — it would be text
    nobody wrote, in a system whose whole job is quoting text somebody did.
    """
    recovered: list[int] = []
    budget = ocr.max_pages

    for index, (raw, has_image) in enumerate(zip(raw_pages, page_has_images, strict=True)):
        page_number = index + 1
        if not _needs_ocr(raw) or not has_image:
            continue
        if budget <= 0:
            logger.warning(
                "ocr: page %d skipped, past the %d-page budget", page_number, ocr.max_pages
            )
            continue

        budget -= 1
        try:
            image = pages[index].to_image(resolution=ocr.dpi).original
            recognized = ocr.recognize(image, page_number=page_number)
        except OCRUnavailableError:
            # The engine itself is broken, so the next page would fail identically.
            raise
        except OCRError as exc:
            # One page failing should not cost the pages that worked; it simply
            # stays reported as having no text.
            logger.warning("ocr: page %d failed (%s)", page_number, type(exc).__name__)
            continue

        if _needs_ocr(recognized):
            logger.info("ocr: page %d yielded no usable text", page_number)
            continue

        raw_pages[index] = recognized
        recovered.append(page_number)

    return tuple(recovered)


def _assemble(
    raw_pages: list[str],
    *,
    has_images: bool = False,
    pages_from_ocr: tuple[int, ...] = (),
    ocr_attempted: bool = False,
    ink: Sequence[_PageInk | None] | None = None,
) -> ParsedDocument:
    """Join per-page text into one offset space, tracking page boundaries.

    `ink` carries per-page character geometry measured against each page's *raw*
    text (M5 slice 3). It defaults to `None` meaning "no geometry", which is what
    `parse_docx` and the tests that call this function with bare strings rely on —
    a `.docx` has no glyph boxes at all, and it must stay possible to assemble a
    document from text alone.
    """
    chunks: list[str] = []
    pages: list[PageSpan] = []
    geometry: list[PageGeometry] = []
    without_text: list[int] = []
    cursor = 0

    for index, raw in enumerate(raw_pages, start=1):
        # Normalize before measuring — see module docstring. U+0000 is stripped
        # here too: a broken ToUnicode map makes extractors emit NUL for glyphs
        # they cannot name, and Postgres refuses NUL in text columns. Removing
        # characters before the spans are measured shifts no offsets. OCR text
        # arrives here as well, so it is held to exactly the same contract.
        page_text = unicodedata.normalize("NFC", raw).replace("\x00", "")

        if _needs_ocr(page_text):
            without_text.append(index)

        if chunks:
            chunks.append(PAGE_SEPARATOR)
            cursor += len(PAGE_SEPARATOR)

        start = cursor
        chunks.append(page_text)
        cursor += len(page_text)
        pages.append(PageSpan(page_number=index, char_start=start, char_end=cursor))

        page_ink = ink[index - 1] if ink is not None and index - 1 < len(ink) else None
        if page_ink is not None:
            rebased = _rebase_ink(page_ink, raw=raw, page_start=start)
            if rebased is not None:
                geometry.append(
                    PageGeometry(
                        page_number=index,
                        width=page_ink.width,
                        height=page_ink.height,
                        runs=rebased,
                    )
                )

    if raw_pages and len(without_text) == len(raw_pages):
        if has_images:
            raise NoTextLayerError(page_count=len(raw_pages), ocr_attempted=ocr_attempted)
        raise EmptyDocumentError(page_count=len(raw_pages))

    return ParsedDocument(
        text="".join(chunks),
        pages=tuple(pages),
        pages_without_text=tuple(without_text),
        pages_from_ocr=pages_from_ocr,
        page_geometry=tuple(geometry),
    )


def _rebase_ink(page_ink: _PageInk, *, raw: str, page_start: int) -> tuple[CharRun, ...] | None:
    """Move one page's geometry from its raw text into the document's offset space.

    Two transforms sit between the two, and they are not the same kind of problem:

    **The NUL strip is exact and is remapped.** A broken ToUnicode map makes an
    extractor emit `\\x00` for glyphs it cannot name, so the characters really do
    move — by up to 11 positions in `resume_broken_tounicode.pdf`, where 8 of 11
    words carry one. Each surviving character's new index is computed and the runs
    are rewritten, splitting where a removal fell inside one.

    **NFC is refused rather than remapped**, and the geometry is dropped if it
    changes anything. Normalization can *combine* characters — `A` plus a combining
    ring is one `Å` afterwards — so there is no index map, and a box drawn from a
    guessed one would be a visual claim nobody could check. Measured: NFC is a no-op
    on every PDF fixture in this repo today, so this costs nothing now and is correct
    the day it does not.
    """
    if unicodedata.normalize("NFC", raw) != raw:
        return None

    positions: list[int | None] = []
    kept = 0
    for character in raw:
        if character == "\x00":
            positions.append(None)
            continue
        positions.append(kept)
        kept += 1

    return geometry_module.shift(geometry_module.remap(page_ink.runs, positions), page_start)
