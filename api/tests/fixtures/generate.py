"""Generate the PDF and DOCX fixtures used by the parser tests.

Run this to (re)create the fixtures, then commit the results:

    python api/tests/fixtures/generate.py

The generated PDFs are committed rather than built during the test run, so CI does
not depend on a Thai-capable system font being installed. This script is the
reproducible record of how they were made.

All content is invented. No real person's resume appears anywhere in this repo.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfdoc, pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

FIXTURE_DIR = Path(__file__).parent
PAGE_W, PAGE_H = A4

# Fonts that ship with Windows and cover Thai. Only needed to *generate* fixtures.
THAI_FONT_CANDIDATES = (
    Path("C:/Windows/Fonts/tahoma.ttf"),
    Path("C:/Windows/Fonts/leelawui.ttf"),
    Path("/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf"),
    Path("/Library/Fonts/Tahoma.ttf"),
)

THAI_FONT_NAME = "ThaiBody"

RESUME_EN = [
    ("H1", "Somchai Jaidee"),
    ("BODY", "Senior Backend Engineer  |  somchai.j@example.com  |  Bangkok, Thailand"),
    ("GAP", ""),
    ("H2", "EXPERIENCE"),
    ("BODY", "Acme Logistics — Backend Engineer (Jan 2021 - Mar 2024)"),
    ("BODY", "  Built payment reconciliation services in Python and PostgreSQL,"),
    ("BODY", "  handling 40,000 transactions per day."),
    ("BODY", "  Cut nightly settlement runtime from 3 hours to 22 minutes."),
    ("GAP", ""),
    ("BODY", "Siam Digital — Junior Developer (Jun 2019 - Dec 2020)"),
    ("BODY", "  Maintained a Django monolith serving 12 internal teams."),
    ("GAP", ""),
    ("H2", "SKILLS"),
    ("BODY", "Python, FastAPI, PostgreSQL, Redis, Docker, pytest"),
    ("GAP", ""),
    ("H2", "EDUCATION"),
    ("BODY", "Chulalongkorn University — B.Eng Computer Engineering (2015 - 2019)"),
]

RESUME_TH = [
    ("H1", "สมชาย ใจดี"),
    ("BODY", "วิศวกรซอฟต์แวร์อาวุโส  |  somchai.j@example.com  |  กรุงเทพมหานคร"),
    ("GAP", ""),
    ("H2", "ประสบการณ์ทำงาน"),
    ("BODY", "บริษัท เอซีเอ็มอี โลจิสติกส์ — วิศวกรซอฟต์แวร์ (ม.ค. 2564 - มี.ค. 2567)"),
    ("BODY", "  ดูแลระบบกระทบยอดการชำระเงินด้วย Python และ PostgreSQL"),
    ("BODY", "  รองรับธุรกรรม 40,000 รายการต่อวัน"),
    ("GAP", ""),
    ("H2", "ทักษะ"),
    ("BODY", "Python, FastAPI, PostgreSQL, Docker, การออกแบบระบบ"),
    ("GAP", ""),
    ("H2", "การศึกษา"),
    ("BODY", "จุฬาลงกรณ์มหาวิทยาลัย — วิศวกรรมศาสตรบัณฑิต สาขาวิศวกรรมคอมพิวเตอร์"),
]


def find_thai_font() -> Path | None:
    return next((p for p in THAI_FONT_CANDIDATES if p.exists()), None)


def register_thai_font() -> str | None:
    font_path = find_thai_font()
    if font_path is None:
        return None
    pdfmetrics.registerFont(TTFont(THAI_FONT_NAME, str(font_path)))
    return THAI_FONT_NAME


def draw_lines(
    c: canvas.Canvas,
    lines: list[tuple[str, str]],
    *,
    body_font: str,
    bold_font: str,
    x: float = 60,
    top: float = PAGE_H - 70,
) -> None:
    y = top
    for kind, text in lines:
        if kind == "GAP":
            y -= 10
            continue
        size = {"H1": 18, "H2": 12, "BODY": 10}[kind]
        font = bold_font if kind in {"H1", "H2"} else body_font
        c.setFont(font, size)
        c.drawString(x, y, text)
        y -= size + 6


def write_single_page(path: Path, lines: list[tuple[str, str]], *, thai: bool) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    if thai:
        font = register_thai_font()
        if font is None:
            raise RuntimeError(
                "No Thai-capable font found; cannot regenerate the Thai fixture. "
                f"Looked in: {', '.join(str(p) for p in THAI_FONT_CANDIDATES)}"
            )
        body_font = bold_font = font
    else:
        body_font, bold_font = "Helvetica", "Helvetica-Bold"
    draw_lines(c, lines, body_font=body_font, bold_font=bold_font)
    c.save()


def write_two_column(path: Path) -> None:
    """Two-column layout — a classic case where naive extraction interleaves text."""
    c = canvas.Canvas(str(path), pagesize=A4)
    left = [
        ("H1", "Nadia Wong"),
        ("H2", "CONTACT"),
        ("BODY", "nadia.w@example.com"),
        ("BODY", "Chiang Mai, Thailand"),
        ("GAP", ""),
        ("H2", "SKILLS"),
        ("BODY", "Go, Kubernetes"),
        ("BODY", "Terraform, gRPC"),
    ]
    right = [
        ("H2", "EXPERIENCE"),
        ("BODY", "Northstar Cloud — SRE"),
        ("BODY", "(Feb 2022 - Present)"),
        ("BODY", "  Ran a 60-node Kubernetes"),
        ("BODY", "  platform for 9 product teams."),
        ("GAP", ""),
        ("BODY", "Highland Systems — DevOps"),
        ("BODY", "(Aug 2020 - Jan 2022)"),
    ]
    draw_lines(c, left, body_font="Helvetica", bold_font="Helvetica-Bold", x=50)
    draw_lines(c, right, body_font="Helvetica", bold_font="Helvetica-Bold", x=300)
    c.save()


def write_two_column_with_header(path: Path) -> None:
    """Two columns under a full-width header, with a full-width footer.

    The shape almost every real two-column resume has, and the one
    `resume_two_column.pdf` is missing: a header line spans the gutter, so a column
    profile taken over the whole page finds nothing. Detection has to cut the page
    into horizontal bands first.
    """
    c = canvas.Canvas(str(path), pagesize=A4)
    draw_lines(
        c,
        [
            ("H1", "Ratana Phongam"),
            ("BODY", "Platform Engineer  |  ratana.p@example.com  |  Khon Kaen, Thailand"),
        ],
        body_font="Helvetica",
        bold_font="Helvetica-Bold",
        x=50,
    )
    left = [
        ("H2", "CONTACT"),
        ("BODY", "ratana.p@example.com"),
        ("BODY", "Khon Kaen, Thailand"),
        ("GAP", ""),
        ("H2", "SKILLS"),
        ("BODY", "Python, Terraform"),
        ("BODY", "Kubernetes, Grafana"),
    ]
    right = [
        ("H2", "EXPERIENCE"),
        ("BODY", "Mekong Payments — Platform (2021 - Present)"),
        ("BODY", "  Ran the deployment platform"),
        ("BODY", "  for 14 engineering teams."),
        ("GAP", ""),
        ("BODY", "Isan Retail — SRE (2018 - 2021)"),
    ]
    body_top = PAGE_H - 150
    draw_lines(c, left, body_font="Helvetica", bold_font="Helvetica-Bold", x=50, top=body_top)
    draw_lines(c, right, body_font="Helvetica", bold_font="Helvetica-Bold", x=300, top=body_top)
    c.setFont("Helvetica", 9)
    c.drawString(50, 60, "References available on request — generated fixture, not a real person")
    c.save()


BLEED = 7.83
"""How far below the origin `resume_two_column_shifted_box.pdf`'s MediaBox starts.

Copied from a real Canva export that `CorruptDocumentError`'d on 2026-08-22 — the
number is arbitrary, but a *plausible* one is worth more than a round one here,
because the bug it reproduces was invisible for exactly as long as every fixture
started at the origin.
"""


class _ShiftedBoxPage(pdfdoc.PDFPage):
    """A page whose MediaBox does not start at (0, 0).

    reportlab writes `[0 0 width height]` and offers no way to ask for anything else,
    but it only fills `MediaBox` in when it is still unset — so a subclass that sets
    it first survives. Patched over `pdfdoc.PDFPage` for one `save()` below.
    """

    def __init__(self) -> None:
        super().__init__()
        self.MediaBox = pdfdoc.PDFArray([0, BLEED, PAGE_W, PAGE_H + BLEED])


def write_two_column_shifted_box(path: Path) -> None:
    """Two columns on a page whose MediaBox starts below the origin.

    **The shape a real resume had and none of these fixtures did.** Design tools
    export a bleed box, so the page runs from `-7.83` to `834.06` rather than from
    `0` to `841.89` — the height is the same and every other reader is unbothered.
    Column detection built its crop boxes out of the page's *lengths*, which are
    still 595.28 x 841.89, so the last band ended 7.83pt below the bottom of the
    page; `pdfplumber.Page.crop` refused it and the whole document came back
    `CorruptDocumentError` — terminal `failed`, not retryable — for a page that reads
    perfectly well. Two columns are what makes it fire, which is to say: the most
    common shape of real resume there is.
    """
    original_page = pdfdoc.PDFPage
    pdfdoc.PDFPage = _ShiftedBoxPage  # type: ignore[misc]
    try:
        c = canvas.Canvas(str(path), pagesize=A4)
        draw_lines(
            c,
            [
                ("H1", "Pimchanok Sirisak"),
                ("BODY", "Data Engineer  |  pimchanok.s@example.com  |  Phuket, Thailand"),
            ],
            body_font="Helvetica",
            bold_font="Helvetica-Bold",
            x=50,
            top=PAGE_H - 90,
        )
        left = [
            ("H2", "CONTACT"),
            ("BODY", "pimchanok.s@example.com"),
            ("BODY", "Phuket, Thailand"),
            ("GAP", ""),
            ("H2", "SKILLS"),
            ("BODY", "Python, Airflow"),
            ("BODY", "dbt, BigQuery"),
        ]
        right = [
            ("H2", "EXPERIENCE"),
            ("BODY", "Andaman Analytics — Data (2021 - Present)"),
            ("BODY", "  Owned the warehouse ingestion"),
            ("BODY", "  for 30 upstream sources."),
            ("GAP", ""),
            ("BODY", "Phuket Freight — Analyst (2018 - 2021)"),
        ]
        body_top = PAGE_H - 170
        draw_lines(c, left, body_font="Helvetica", bold_font="Helvetica-Bold", x=50, top=body_top)
        draw_lines(c, right, body_font="Helvetica", bold_font="Helvetica-Bold", x=300, top=body_top)
        c.save()
    finally:
        pdfdoc.PDFPage = original_page  # type: ignore[misc]


def write_right_aligned_dates(path: Path) -> None:
    """One column, with the dates pushed out to the right margin.

    The false positive column detection has to refuse. Every role line leaves a wide
    empty strip between the job title and its date, which is the shape of a gutter;
    what makes it *not* one is that the bullets underneath run the full width of the
    page. A page like this must keep parsing exactly as it did before.
    """
    c = canvas.Canvas(str(path), pagesize=A4)
    rows = [
        ("Wanida Chaiyo", "", 18),
        ("Backend Engineer  |  wanida.c@example.com  |  Phuket, Thailand", "", 10),
        ("", "", 0),
        ("EXPERIENCE", "", 12),
        ("Andaman Software — Backend Engineer", "2022 - Present", 10),
        ("  Built the booking API in FastAPI and PostgreSQL, serving 11 partner", "", 10),
        ("  agencies and roughly 8,000 reservations a day across three regions.", "", 10),
        ("  Cut p95 checkout latency from 900ms to 210ms by batching the writes.", "", 10),
        ("Similan Tech — Developer", "2019 - 2022", 10),
        ("  Maintained a PHP storefront and migrated its billing to a queue-based", "", 10),
        ("  worker, which removed the nightly maintenance window entirely.", "", 10),
        ("", "", 0),
        ("EDUCATION", "", 12),
        ("Prince of Songkla University — B.Sc Computer Science", "2015 - 2019", 10),
    ]
    y = PAGE_H - 70
    for text, date, size in rows:
        if not size:
            y -= 10
            continue
        c.setFont("Helvetica-Bold" if size > 10 else "Helvetica", size)
        c.drawString(50, y, text)
        if date:
            c.setFont("Helvetica", 10)
            c.drawRightString(PAGE_W - 50, y, date)
        y -= size + 6
    c.save()


def write_multipage(path: Path, pages: int = 3) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    for page in range(1, pages + 1):
        lines: list[tuple[str, str]] = [("H1", f"Project Portfolio — Page {page}")]
        for item in range(1, 6):
            marker = f"P{page}I{item}"
            lines.append(("BODY", f"Page {page} project {item}: distinctive marker {marker}."))
        draw_lines(c, lines, body_font="Helvetica", bold_font="Helvetica-Bold")
        c.showPage()
    c.save()


def _render_text_image(lines: list[str], *, width: int = 1240, height: int = 1754) -> Image.Image:
    """Rasterize text so the resulting PDF page has no text layer at all."""
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    font_path = find_thai_font()
    try:
        font = ImageFont.truetype(str(font_path), 28) if font_path else ImageFont.load_default()
    except OSError:
        font = ImageFont.load_default()

    y = 90
    for line in lines:
        draw.text((80, y), line, fill="black", font=font)
        y += 46

    # Faint speckle so it reads as a scan rather than a clean render.
    for offset in range(0, width, 37):
        draw.point((offset, (offset * 7) % height), fill="gray")
    return image


def write_scanned(path: Path) -> None:
    """A resume that exists only as an image — must raise NoTextLayerError in M1
    and is the OCR fixture for M2."""
    image = _render_text_image(
        [
            "Somchai Jaidee",
            "Senior Backend Engineer",
            "",
            "EXPERIENCE",
            "Acme Logistics - Backend Engineer (Jan 2021 - Mar 2024)",
            "Built payment reconciliation services in Python.",
            "",
            "ทักษะ: Python, FastAPI, PostgreSQL",
        ]
    )
    c = canvas.Canvas(str(path), pagesize=A4)
    c.drawImage(ImageReader(image), 0, 0, width=PAGE_W, height=PAGE_H)
    c.save()


def write_mixed_scan(path: Path) -> None:
    """Page 1 has a text layer, page 2 does not — partial-scan handling."""
    c = canvas.Canvas(str(path), pagesize=A4)
    draw_lines(
        c,
        [
            ("H1", "Preecha Boonmee"),
            ("H2", "SUMMARY"),
            ("BODY", "Data platform engineer with 6 years building ETL pipelines."),
        ],
        body_font="Helvetica",
        bold_font="Helvetica-Bold",
    )
    c.showPage()
    image = _render_text_image(["EXPERIENCE", "Riverbank Analytics - Data Engineer"])
    c.drawImage(ImageReader(image), 0, 0, width=PAGE_W, height=PAGE_H)
    c.save()


# Lines for the damaged fixture. Deliberately ASCII and free of parentheses and
# backslashes, which would need escaping inside a PDF string literal.
BROKEN_FONT_LINES = [
    "Somchai Jaidee",
    "Backend Engineer at Acme Logistics",
    "Skills: Python, FastAPI, PostgreSQL",
]

# Glyphs whose ToUnicode entry will point at U+0000 instead of the real character.
BROKEN_FONT_GLYPHS = frozenset("ai")


def _broken_tounicode_cmap(text: str, broken: frozenset[str]) -> str:
    """A ToUnicode CMap that maps some glyphs to U+0000.

    This is the damage, stated exactly: the map exists and is well-formed, so
    nothing refuses to parse — it simply says that certain glyphs mean nothing.
    """
    codes = sorted({ord(character) for character in text if character.strip()})
    entries = "\n".join(
        f"<{code:02X}> <{0 if chr(code) in broken else code:04X}>" for code in codes
    )
    return f"""/CIDInit /ProcSet findresource begin
12 dict begin
begincmap
/CMapName /BrokenToUnicode def
/CMapType 2 def
1 begincodespacerange
<00> <FF>
endcodespacerange
{len(codes)} beginbfchar
{entries}
endbfchar
endcmap
CMapName currentdict /CMap defineresource pop
end
end"""


def write_broken_tounicode(path: Path) -> None:
    """A PDF whose font maps some glyphs to U+0000 — the shape of the §11 incident.

    A real Thai resume template (a designer-tool export) had eight glyphs whose
    embedded font carried no usable ToUnicode mapping, and pdfplumber returned
    U+0000 for each. Postgres refuses NUL in a text column, so the pipeline
    succeeded completely and then failed at the commit. SQLite stores NUL happily,
    which is why the whole suite was blind to it.

    That is now stripped in `_assemble` before page spans are measured, and pinned
    at the seam by `TestControlCharacters`. This fixture covers *the road to it*:
    a real PDF, parsed by the real parser, actually producing NUL.

    Written by hand rather than with reportlab, because the damage has to be in the
    font's ToUnicode map and reportlab only ever writes correct ones. Removing the
    map entirely — the obvious approach — turns out to produce `(cid:1)(cid:2)...`
    placeholders instead of NUL, which is a different defect. Hand-writing it is
    also why this fixture needs no Thai font to regenerate.
    """
    text = "".join(BROKEN_FONT_LINES)
    stream = "BT /F1 14 Tf 60 760 Td " + " ".join(
        f"({line}) Tj 0 -22 Td" for line in BROKEN_FONT_LINES
    )
    content = f"{stream} ET".encode("latin-1")
    tounicode = _broken_tounicode_cmap(text, BROKEN_FONT_GLYPHS).encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /ToUnicode 6 0 R >>",
        b"<< /Length %d >>\nstream\n" % len(content) + content + b"\nendstream",
        b"<< /Length %d >>\nstream\n" % len(tounicode) + tounicode + b"\nendstream",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % number + body + b"\nendobj\n"

    xref_at = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        xref_at,
    )
    path.write_bytes(bytes(out))


def write_empty(path: Path) -> None:
    """A structurally valid PDF with a blank page."""
    c = canvas.Canvas(str(path), pagesize=A4)
    c.showPage()
    c.save()


def write_not_a_pdf(path: Path) -> None:
    path.write_bytes(b"This is plainly not a PDF file.\n" * 4)


def write_docx(path: Path) -> None:
    """A Word resume whose skills sit in a table.

    The table is the point: resumes routinely use one for layout, and a parser that
    reads only `document.paragraphs` drops it silently — which looks like a model
    that missed the skills rather than a parser that never saw them. Thai is here
    for the same reason it is in the PDF fixtures.
    """
    document = Document()
    document.add_paragraph("Kanya Sriwong")
    document.add_paragraph("วิศวกรข้อมูล  |  kanya.s@example.com  |  เชียงใหม่")

    document.add_paragraph("EXPERIENCE")
    document.add_paragraph("Lanna Data — Data Engineer (Feb 2022 - Present)")
    document.add_paragraph("  สร้างไปป์ไลน์ ETL ด้วย Python และ Airflow")

    document.add_paragraph("SKILLS")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "6 years"
    table.cell(1, 0).text = "Airflow, dbt"
    table.cell(1, 1).text = "4 years"

    document.add_paragraph("การศึกษา")
    document.add_paragraph("มหาวิทยาลัยเชียงใหม่ — วิศวกรรมคอมพิวเตอร์")
    document.save(str(path))


def write_empty_docx(path: Path) -> None:
    """A structurally valid .docx with nothing in it."""
    Document().save(str(path))


def main() -> int:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    write_single_page(FIXTURE_DIR / "resume_en.pdf", RESUME_EN, thai=False)
    write_single_page(FIXTURE_DIR / "resume_th.pdf", RESUME_TH, thai=True)
    write_two_column(FIXTURE_DIR / "resume_two_column.pdf")
    write_two_column_with_header(FIXTURE_DIR / "resume_two_column_header.pdf")
    write_two_column_shifted_box(FIXTURE_DIR / "resume_two_column_shifted_box.pdf")
    write_right_aligned_dates(FIXTURE_DIR / "resume_right_aligned_dates.pdf")
    write_multipage(FIXTURE_DIR / "resume_multipage.pdf")
    write_scanned(FIXTURE_DIR / "resume_scanned.pdf")
    write_mixed_scan(FIXTURE_DIR / "resume_mixed_scan.pdf")
    write_broken_tounicode(FIXTURE_DIR / "resume_broken_tounicode.pdf")
    write_empty(FIXTURE_DIR / "empty.pdf")
    write_not_a_pdf(FIXTURE_DIR / "not_a_pdf.pdf")
    write_docx(FIXTURE_DIR / "resume_th.docx")
    write_empty_docx(FIXTURE_DIR / "empty.docx")

    for fixture in sorted(FIXTURE_DIR.glob("*.pdf")) + sorted(FIXTURE_DIR.glob("*.docx")):
        print(f"  {fixture.name:28} {fixture.stat().st_size:>8,} bytes")
    return 0


if __name__ == "__main__":
    sys.exit(main())
